import io
import os
import uuid
import logging
import docx
import PyPDF2
import numpy as np
from openai import OpenAI
from dotenv import load_dotenv
from pydantic import BaseModel
from typing import List, Optional
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from sklearn.metrics.pairwise import cosine_similarity
from fastapi.responses import JSONResponse, FileResponse
from fastapi import FastAPI, HTTPException, UploadFile, File, Form



logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()

app = FastAPI(title="Document Q&A API", version="1.0.0")

app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
def read_root():
    return FileResponse(os.path.join("static", "index.html"))

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize OpenAI client
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    logger.error("OPENAI_API_KEY not found in environment variables!")
    raise ValueError("OPENAI_API_KEY not found. Please set it in your .env file")

client = OpenAI(api_key=api_key)
logger.info("OpenAI client initialized successfully")

documents_store = {}
vectors_store = {}


class Document(BaseModel):
    title: str
    content: str

class DocumentList(BaseModel):
    documents: List[Document]

class Query(BaseModel):
    question: str
    top_k: Optional[int] = 3

class Answer(BaseModel):
    answer: str
    sources: List[dict]

class DocumentResponse(BaseModel):
    id: str
    title: str
    chunks_count: int




#---------------------------------------- FUNCTIONS ----------------------------------------#
    
def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error reading PDF file: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    try:
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error reading DOCX file: {str(e)}")

def extract_text_from_txt(file_content: bytes) -> str:
    try:
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                return file_content.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise ValueError("Unable to decode file with common encodings")
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error reading TXT file: {str(e)}")

def extract_text_from_file(filename: str, file_content: bytes) -> str:
    extension = filename.lower().split('.')[-1]
    
    if extension == 'pdf':
        return extract_text_from_pdf(file_content)
    elif extension == 'docx':
        return extract_text_from_docx(file_content)
    elif extension in ['txt', 'md']:
        return extract_text_from_txt(file_content)
    else:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type: {extension}. Supported types: pdf, docx, txt, md"
        )

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    if not text or not text.strip():
        return [text]
    
    words = text.split()
    if len(words) <= chunk_size:
        return [text]
    
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
    return chunks if chunks else [text]

def get_embedding(text: str) -> List[float]:
    try:
        logger.info(f"Getting embedding for text of length: {len(text)}")
        response = client.embeddings.create(
            model="text-embedding-3-small",
            input=text
        )
        logger.info("Embedding created successfully")
        return response.data[0].embedding
    except Exception as e:
        logger.error(f"Error getting embedding: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting embedding: {str(e)}")

def generate_answer(question: str, context: str) -> str:
    try:
        logger.info(f"Generating answer for question: {question[:50]}...")
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that answers questions based on the provided context. If the answer is not in the context, say so."},
                {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {question}\n\nAnswer:"}
            ],
            temperature=0.7,
            max_tokens=500
        )
        logger.info("Answer generated successfully")
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"Error generating answer: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating answer: {str(e)}")

def index_document_content(title: str, content: str) -> DocumentResponse:
    if not content or not content.strip():
        raise HTTPException(status_code=400, detail="Document content cannot be empty")
    
    doc_id = str(uuid.uuid4())
    
    chunks = chunk_text(content)
    logger.info(f"Created {len(chunks)} chunks for document")
    
    documents_store[doc_id] = {
        "id": doc_id,
        "title": title,
        "content": content,
        "chunks": chunks
    }
    
    for idx, chunk in enumerate(chunks):
        logger.info(f"Creating embedding for chunk {idx + 1}/{len(chunks)}")
        embedding = get_embedding(chunk)
        chunk_id = f"{doc_id}_chunk_{idx}"
        
        vectors_store[chunk_id] = {
            "doc_id": doc_id,
            "chunk_id": chunk_id,
            "title": title,
            "chunk": chunk,
            "embedding": embedding,
            "chunk_index": idx
        }
    
    logger.info(f"Successfully indexed document: {title}")
    
    return DocumentResponse(
        id=doc_id,
        title=title,
        chunks_count=len(chunks)
    )

#---------------------------------------- API ENDPOINTS ----------------------------------------#

@app.post("/documents", response_model=List[DocumentResponse])
async def index_documents(doc_list: DocumentList):
    try:
        logger.info(f"Received request to index {len(doc_list.documents)} documents")
        responses = []
        
        for doc in doc_list.documents:
            logger.info(f"Processing document: {doc.title}")
            response = index_document_content(doc.title, doc.content)
            responses.append(response)
        
        return responses
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in index_documents: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

@app.post("/upload", response_model=List[DocumentResponse])
async def upload_files(files: List[UploadFile] = File(...)):
    try:
        logger.info(f"Received {len(files)} files for upload")
        responses = []
        
        for file in files:
            logger.info(f"Processing file: {file.filename}")
            
            file_content = await file.read()
            
            text_content = extract_text_from_file(file.filename, file_content)
            
            if not text_content or not text_content.strip():
                logger.warning(f"No text extracted from file: {file.filename}")
                raise HTTPException(
                    status_code=400, 
                    detail=f"No text could be extracted from {file.filename}"
                )
            
            response = index_document_content(file.filename, text_content)
            responses.append(response)
        
        return responses
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in upload_files: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

@app.post("/query", response_model=Answer)
async def query_documents(query: Query):
    try:
        logger.info(f"Received query: {query.question}")
        
        if not vectors_store:
            raise HTTPException(status_code=400, detail="No documents indexed yet")
        
        query_embedding = get_embedding(query.question)
        
        similarities = []
        for chunk_id, vector_data in vectors_store.items():
            similarity = cosine_similarity(
                [query_embedding],
                [vector_data["embedding"]]
            )[0][0]
            
            similarities.append({
                "chunk_id": chunk_id,
                "doc_id": vector_data["doc_id"],
                "title": vector_data["title"],
                "chunk": vector_data["chunk"],
                "similarity": float(similarity)
            })
        
        similarities.sort(key=lambda x: x["similarity"], reverse=True)
        top_results = similarities[:query.top_k]
        
        logger.info(f"Found {len(top_results)} relevant chunks")
        
        context = "\n\n".join([f"[{r['title']}]\n{r['chunk']}" for r in top_results])
        
        answer_text = generate_answer(query.question, context)
        
        sources = [
            {
                "doc_id": r["doc_id"],
                "title": r["title"],
                "chunk": r["chunk"][:200] + "..." if len(r["chunk"]) > 200 else r["chunk"],
                "similarity_score": r["similarity"]
            }
            for r in top_results
        ]
        
        return Answer(answer=answer_text, sources=sources)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in query_documents: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

@app.get("/documents", response_model=List[DocumentResponse])
async def list_documents():
    try:
        return [
            DocumentResponse(
                id=doc_id,
                title=doc_data["title"],
                chunks_count=len(doc_data["chunks"])
            )
            for doc_id, doc_data in documents_store.items()
        ]
    except Exception as e:
        logger.error(f"Error listing documents: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error listing documents: {str(e)}")

@app.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    try:
        if doc_id not in documents_store:
            raise HTTPException(status_code=404, detail="Document not found")
        
        del documents_store[doc_id]
        
        chunks_to_remove = [
            chunk_id for chunk_id, vector_data in vectors_store.items()
            if vector_data["doc_id"] == doc_id
        ]
        
        for chunk_id in chunks_to_remove:
            del vectors_store[chunk_id]
        
        logger.info(f"Deleted document {doc_id} with {len(chunks_to_remove)} chunks")
        
        return {"message": f"Document {doc_id} deleted successfully", "chunks_removed": len(chunks_to_remove)}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

@app.get("/health")
async def root():
    return {
        "status": "running",
        "api_key_configured": bool(api_key),
        "documents_count": len(documents_store),
        "chunks_count": len(vectors_store)
    }

@app.get("/test-openai")
async def test_openai():
    """Test OpenAI connection"""
    try:
        response = client.embeddings.create(
            model="text-embedding-3-small",
            input="test"
        )
        return {"status": "success", "message": "OpenAI API is working correctly"}
    except Exception as e:
        logger.error(f"OpenAI test failed: {str(e)}")
        return {"status": "error", "message": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)